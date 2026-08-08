"""Génère les pièces sources du cas fictif en format Word (.docx) éditable — pour relecture/copier-coller admin."""
import os
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor

OUT = "/app/frontend/public/cas-demonstration/word"
os.makedirs(OUT, exist_ok=True)
RED = RGBColor(0xB0, 0x14, 0x14)


def make_docx(filename, title, sender, blocks):
    d = Document()
    banner = d.add_paragraph()
    r = banner.add_run("CAS FICTIF DE DÉMONSTRATION — Aucune personne réelle. Document créé pour l'évaluation de Dossier Express IA.")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RED
    if sender:
        p = d.add_paragraph()
        rs = p.add_run(sender)
        rs.font.size = Pt(9)
    t = d.add_paragraph()
    rt = t.add_run(title)
    rt.bold = True
    rt.font.size = Pt(13)
    for bold, text in blocks:
        p = d.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10.5)
    d.save(f"{OUT}/{filename}")
    print("OK", filename)


def blocks_from_lines(lines):
    out = []
    for line in lines:
        if line.startswith("**"):
            out.append((True, line[2:]))
        else:
            out.append((False, line))
    return out


make_docx("1-certificat-medical-initial.docx",
    "CERTIFICAT MÉDICAL INITIAL — Maladie professionnelle (Cerfa fictif)",
    "Dr Paul EXEMPLE — Cabinet médical fictif\n12 rue de la Démonstration, 99000 Villefictive",
    blocks_from_lines([
        "Je soussigné, Dr Paul EXEMPLE, médecin généraliste (RPPS fictif 99999999901), certifie avoir examiné ce jour :",
        "**Mme DEMONSTRATION Claire, née le 12/03/1981 (identité fictive)",
        "Profession déclarée : cadre de proximité en établissement médico-social.",
        "**Constatations :",
        "Syndrome dépressif caractérisé d'intensité sévère : humeur triste permanente, anhédonie, troubles du sommeil majeurs, ruminations anxieuses centrées sur le travail, perte de poids (6 kg en 3 mois), idéation suicidaire passive sans scénario. Symptomatologie apparue progressivement dans un contexte décrit de surcharge professionnelle durable, d'objectifs contradictoires et d'isolement managérial.",
        "**Diagnostic évoqué : épisode dépressif caractérisé sévère, sans symptôme psychotique, en lien évoqué avec l'activité professionnelle (maladie hors tableau).",
        "Date de première constatation médicale : 18/02/2026.",
        "Arrêt de travail initial prescrit. Orientation vers suivi psychiatrique spécialisé.",
        "Fait à Villefictive, le 18/02/2026 — Dr Paul EXEMPLE (signature fictive)"]))

make_docx("2-notification-refus-cpam.docx",
    "NOTIFICATION DE DÉCISION — Refus de reconnaissance de maladie professionnelle",
    "CPAM FICTIVE DE LA DÉMONSTRATION\nService Risques Professionnels — 99000 Villefictive",
    blocks_from_lines([
        "Dossier n° MP-FICTIF-2026-00427",
        "Assurée : Mme DEMONSTRATION Claire (NIR fictif 2 81 03 99 999 999)",
        "Madame,",
        "Vous avez déclaré le 24/02/2026 une maladie « épisode dépressif sévère » que vous estimez d'origine professionnelle.",
        "**Après instruction, votre maladie ne figure dans aucun tableau de maladies professionnelles du régime général.",
        "Conformément à l'article L.461-1 du code de la sécurité sociale, la reconnaissance au titre du système complémentaire nécessite que l'incapacité permanente prévisible soit au moins égale à 25 % et que le lien direct et essentiel avec le travail habituel soit établi par le Comité Régional de Reconnaissance des Maladies Professionnelles (CRRMP).",
        "Le médecin-conseil a évalué votre taux d'incapacité permanente prévisible à 20 %, taux inférieur au seuil requis. En conséquence, votre dossier n'a pas été transmis au CRRMP et votre demande est REJETÉE.",
        "**Voies de recours : vous disposez de deux mois pour saisir la Commission de Recours Amiable (CRA) de notre organisme, et/ou de contester l'évaluation médicale.",
        "Villefictive, le 12/05/2026 — CPAM fictive de la Démonstration"]))

make_docx("3-compte-rendu-psychiatrique.docx",
    "COMPTE RENDU DE SUIVI PSYCHIATRIQUE",
    "Dr Anne SPECIMEN — Psychiatre (cabinet fictif)\n4 avenue de l'Exemple, 99000 Villefictive",
    blocks_from_lines([
        "Patiente : Mme DEMONSTRATION Claire, 45 ans (identité fictive)",
        "Suivie depuis le 03/03/2026 — 6 consultations à ce jour.",
        "**Anamnèse : cadre de proximité depuis 2015, décrit depuis 2023 une dégradation continue des conditions d'exercice : sous-effectif chronique, injonctions contradictoires, astreintes non compensées, alerte adressée à la direction en 09/2025 restée sans réponse. Pas d'antécédent psychiatrique personnel ou familial. Pas de facteur de vulnérabilité extra-professionnel identifié à ce jour.",
        "**Clinique actuelle : épisode dépressif caractérisé d'intensité sévère (PHQ-9 : 21/27). Ruminations exclusivement centrées sur la sphère professionnelle, reviviscences anxieuses à l'évocation du lieu de travail, évitement. Amélioration partielle depuis l'éloignement du poste (arrêt de travail continu).",
        "**Traitement : sertraline 100 mg/j (fictif), psychothérapie hebdomadaire.",
        "**Évolution : stabilisation lente ; incapacité durable prévisible en l'état. L'imputabilité professionnelle apparaît cliniquement prépondérante.",
        "Villefictive, le 20/06/2026 — Dr Anne SPECIMEN, psychiatre (RPPS fictif 99999999902)"]))

make_docx("4-arret-travail-VERSION-LISIBLE.docx",
    "AVIS D'ARRÊT DE TRAVAIL (prolongation) — [NB : la version PDF de cette pièce est volontairement floue pour tester le contrôle qualité]",
    "",
    blocks_from_lines([
        "Cerfa fictif n° 10170*07",
        "Patiente : Mme DEMONSTRATION Claire (identité fictive)",
        "Prolongation du 15/04/2026 au 15/06/2026 inclus",
        "Motif : épisode dépressif sévère — lien professionnel évoqué",
        "Sorties autorisées : oui, en dehors des heures de présence obligatoire",
        "Dr Paul EXEMPLE — 18/04/2026 (signature fictive illisible)"]))

make_docx("5-elements-contexte-professionnel.docx",
    "ATTESTATION ET ÉLÉMENTS DE CONTEXTE PROFESSIONNEL",
    "Résidence de la Démonstration (établissement fictif)\nInstance représentative du personnel",
    blocks_from_lines([
        "Je soussignée, Mme Sophie MODELE, déléguée du personnel (fictive) de l'établissement « Résidence de la Démonstration » (structure fictive), atteste des faits suivants :",
        "**1. Depuis janvier 2023, le service encadré par Mme DEMONSTRATION fonctionne en sous-effectif permanent (2 postes vacants non remplacés sur 8).",
        "**2. Un signalement écrit sur la surcharge de l'encadrement a été transmis à la direction le 14/09/2025 (copie jointe au CSE fictif). Aucune mesure n'a suivi.",
        "**3. Les astreintes de week-end ont doublé entre 2024 et 2026 sans compensation.",
        "**4. Deux autres salariés du même service sont en arrêt pour des motifs similaires.",
        "Fiche de poste jointe : encadrement d'équipe, continuité de service, gestion des situations de crise, disponibilité permanente téléphonique.",
        "Fait pour servir et valoir ce que de droit (document fictif de démonstration).",
        "Villefictive, le 02/04/2026 — Mme Sophie MODELE"]))

make_docx("6-courrier-medecin-conseil.docx",
    "COURRIER DU SERVICE MÉDICAL — Évaluation du taux d'incapacité prévisible",
    "SERVICE MÉDICAL — CPAM fictive de la Démonstration",
    blocks_from_lines([
        "Réf. dossier MP-FICTIF-2026-00427 — Mme DEMONSTRATION Claire",
        "Madame,",
        "À la suite de l'examen du 28/04/2026, le service médical retient :",
        "**- épisode dépressif caractérisé, d'intensité jugée « modérée à sévère » à l'examen ;",
        "**- retentissement socioprofessionnel présent mais estimé « partiellement réversible » ;",
        "**- taux d'incapacité permanente PRÉVISIBLE évalué à 20 %.",
        "Ce taux étant inférieur au seuil de 25 % prévu à l'article R.461-8 du code de la sécurité sociale, les conditions de saisine du CRRMP ne sont pas réunies.",
        "Il vous est loisible de produire tout élément médical nouveau de nature à modifier cette évaluation (expertise psychiatrique détaillée, échelles d'évaluation, retentissement fonctionnel documenté).",
        "Villefictive, le 05/05/2026 — Dr Marc TYPE, médecin-conseil (fictif)"]))

with zipfile.ZipFile(f"{OUT}/cas-demonstration-word.zip", "w") as z:
    for f in sorted(os.listdir(OUT)):
        if f.endswith(".docx"):
            z.write(f"{OUT}/{f}", f)
print("ZIP OK —", sorted(os.listdir(OUT)))
